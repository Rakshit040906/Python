# Create an editable DOCX based on the inspection report structure
from docx import Document
from docx.shared import Inches

doc = Document()

# Title
doc.add_heading('FALABELLA INDIA', 0)
doc.add_heading('SHOES FINAL INSPECTION REPORT', 1)

# Basic Info Table
table = doc.add_table(rows=6, cols=4)
fields = [
    ("Style No", "", "Vendor", ""),
    ("Description", "", "Factory", ""),
    ("Line", "", "Color", ""),
    ("Department", "", "Order Qty", ""),
    ("Ship Date", "", "Ship Qty", ""),
    ("TTL CTNS", "", "Country", "")
]

for i, row in enumerate(fields):
    for j in range(4):
        table.rows[i].cells[j].text = row[j]

doc.add_paragraph("\n")

# Section function
def add_section(title, items):
    doc.add_heading(title, 2)
    t = doc.add_table(rows=len(items)+1, cols=3)
    t.rows[0].cells[0].text = "Check Points"
    t.rows[0].cells[1].text = "Major"
    t.rows[0].cells[2].text = "Minor"
    
    for i, item in enumerate(items):
        t.rows[i+1].cells[0].text = item

# Sections
add_section("1. Packing Accuracy", [
    "Poor carton quality",
    "Incorrect quantity per carton",
    "Wrong content/assortment",
    "Incorrect labeling/marks",
    "Incorrect packing method"
])

add_section("2. Barcodes", [
    "Barcode not readable",
    "Incorrect barcode",
    "Wrong barcode position"
])

add_section("3. Material Defects", [
    "Damaged material",
    "Wrinkles/creases",
    "Incorrect construction",
    "Mismatch of design",
    "Color variation"
])

add_section("4. Construction & Stitches", [
    "Staples/nails",
    "Uneven heel height",
    "Uneven sole",
    "Loose threads",
    "Improper stitching"
])

add_section("5. Product Logo, Labeling & Markings", [
    "Missing label",
    "Incorrect labeling info",
    "Print defects",
    "Color mismatch"
])

add_section("6. Finish", [
    "Sharp edges",
    "Threads untrimmed",
    "Creases/marks",
    "Offensive odor"
])

add_section("7. Cleanliness", [
    "Spots/stains",
    "Glue marks",
    "Other marks"
])

add_section("8. Measurement/Fit", [
    "Out of tolerance",
    "Fit issue"
])

add_section("9. Other Tests", [
    "Velcro failure",
    "Pull test failed",
    "Color fastness failed",
    "Waterproof failed"
])

add_section("10. Others", [
    "Children footwear test",
    "Other issues"
])

doc.add_paragraph("\nInspection Result: PASS")

file_path = "/mnt/data/Shoes_Final_Inspection_Report.docx"
doc.save(file_path)

file_path  
 